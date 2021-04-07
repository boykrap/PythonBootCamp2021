import wikipedia
wikipedia.set_lang('th')

#python to docx..........................

from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    data = wikipedia.summary(keyword)

    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document()
    doc.add_heading(keyword, 0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('complete')


#GUI..............................
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.geometry('500x300')
GUI.title('Wikipedia')

#config...........................
FONT1 = ('Angsana new',15)
FONT2 = ('Angsana new',25)

L = ttk.Label(GUI,text='ค้นหาบทความ',font=FONT2)
L.pack()

#Entry................................
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)


#button Search--------------------------
def Search():
    keyword = v_search.get()
    try:
        lang = v_radio.get()
        Wiki(keyword,lang)
        messagebox.showinfo('Record completed','Sucess OK!!!')
    except:
        messagebox.showwarning('Keyword error','Entry again')
        
    #resultSearch = wikipedia.search(keyword)
    #result = wikipedia.summary(keyword)
    #print(resultSearch)
    #print(result)

B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10)

#Select lang--------------------------
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar()
RB1 = ttk.Radiobutton(F1,text='Thai',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='English',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='China',variable=v_radio,value='zh')
RB1.invoke()

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)






GUI.mainloop()

