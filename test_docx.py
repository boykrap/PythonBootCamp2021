from docx import Document
import wikipedia

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    data = wikipedia.summary(keyword)

    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document()
    doc.add_heading(keyword, 0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('complete')

try:
    Wiki('ereerere','en')
except:
    print('Error')

#Wiki('ประเทศจีน','lo')
#Wiki('ประเทศญี่ปุ่น','en')
#Wiki('ประเทศเวียดนาม','zh')
